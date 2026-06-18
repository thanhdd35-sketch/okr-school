from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import unicodedata
from urllib.parse import quote
from datetime import date
from database import supabase
from auth import chi_giao_vien

router = APIRouter()

TEN_TRUONG = "TRƯỜNG THCS VÀ THPT FPT"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "logo-fpt.jpg")

# Màu chữ
ORANGE = RGBColor(0xF2, 0x8C, 0x28)
NAVY   = RGBColor(0x1F, 0x4E, 0x79)
DNAVY  = RGBColor(0x17, 0x36, 0x5D)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1A, 0x1A, 0x1A)

# Màu nền ô (hex, không dấu #)
F_INFO   = "F2F4F7"
F_ORANGE = "FFF7ED"
F_BLUE   = "EAF2F8"
F_GREEN  = "E2F0D9"
F_CORAL  = "FCE4D6"
F_NAVY   = "1F4E79"
F_DNAVY  = "17365D"
F_LGRAY  = "F7F7F7"

TRANG_THAI_TD = {
    "xuat_sac":   "Xuất sắc (≥ 100%)",
    "tot":        "Tốt (≥ 80%)",
    "dung_huong": "Đúng hướng (≥ 60%)",
    "can_chu_y":  "Cần chú ý (≥ 40%)",
    "chech_huong":"Chệch hướng (< 40%)",
}
TRANG_THAI_MT = {
    "da_duyet": "Đã duyệt", "cho_duyet": "Chờ duyệt",
    "yeu_cau_sua": "Cần sửa", "nhap": "Nháp",
}
FILL_TRANG_THAI = {
    "da_duyet": F_GREEN, "cho_duyet": F_BLUE,
    "yeu_cau_sua": F_CORAL, "nhap": F_LGRAY,
}


def _content_disposition(ten_file: str) -> str:
    ascii_fallback = unicodedata.normalize("NFKD", ten_file).encode("ascii", "ignore").decode("ascii")
    if not ascii_fallback.strip("_. "):
        ascii_fallback = "BaoCao.docx"
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(ten_file)}"


def _shade(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell(cell, lines, *, fill=None, valign="center"):
    """lines: list of (text, size, bold, color). Mỗi dòng = 1 run trong 1 paragraph."""
    if fill:
        _shade(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if valign == "center" else WD_ALIGN_VERTICAL.TOP
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, (text, size, bold, color) in enumerate(lines):
        run = p.add_run(("" if i == 0 else "\n") + str(text))
        run.font.size = Pt(size)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color


def _heading(doc, text, size, color, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def _no_border(table):
    table.style = "Table Grid"


def tao_phieu_hoc_sinh(doc: Document, hs: dict, muc_tieu_list: list,
                       krs_by_mt: dict, danh_gia: dict, ten_ky: str, ten_gv: str):
    # ---- Logo ----
    if os.path.exists(LOGO_PATH):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(LOGO_PATH, width=Cm(3.8))
        except Exception:
            pass

    # ---- Tiêu đề ----
    p_truong = doc.add_paragraph()
    p_truong.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_truong.add_run(TEN_TRUONG)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ORANGE

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("PHIẾU ĐÁNH GIÁ MỤC TIÊU OKR")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run(ten_ky)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = DNAVY
    # gạch chân navy dưới tiêu đề
    pPr = p_sub._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "6"); bottom.set(qn("w:color"), "1F4E79")
    pbdr.append(bottom); pPr.append(pbdr)

    # ---- Bảng thông tin ----
    tds = [mt.get("tien_do_tong") or 0 for mt in muc_tieu_list if mt.get("trang_thai") == "da_duyet"]
    pct_ky = round(sum(tds) / len(tds), 1) if tds else 0
    so_da_duyet = sum(1 for mt in muc_tieu_list if mt.get("trang_thai") == "da_duyet")
    so_can_sua = sum(1 for mt in muc_tieu_list if mt.get("trang_thai") == "yeu_cau_sua")
    tinh_trang = "Đang theo dõi" if muc_tieu_list else "Chưa nộp mục tiêu"

    info = doc.add_table(rows=2, cols=4); _no_border(info)
    lbl = lambda t: (t, 8.5, True, GRAY)
    val = lambda t: (str(t), 11, False, DARK)
    r0 = info.rows[0].cells
    _set_cell(r0[0], [lbl("HỌ TÊN"), val(hs.get("ho_ten", ""))], fill=F_INFO)
    _set_cell(r0[1], [lbl("LỚP"), val(hs.get("ten_lop", ""))], fill=F_INFO)
    _set_cell(r0[2], [lbl("NGÀY XUẤT"), val(date.today().strftime("%d/%m/%Y"))], fill=F_INFO)
    _set_cell(r0[3], [lbl("KỲ ĐÁNH GIÁ"), val(ten_ky)], fill=F_INFO)
    r1 = info.rows[1].cells
    _set_cell(r1[0], [lbl("GVCN"), val(ten_gv)])
    _set_cell(r1[1], [lbl("TIẾN ĐỘ TỔNG KỲ"), val(f"{pct_ky}%")])
    _set_cell(r1[2], [lbl("SỐ MỤC TIÊU ĐÃ NỘP"), val(len(muc_tieu_list))])
    _set_cell(r1[3], [lbl("TÌNH TRẠNG"), val(tinh_trang)])

    # ---- Tổng quan kết quả ----
    _heading(doc, "TỔNG QUAN KẾT QUẢ", 13, NAVY)
    tq = doc.add_table(rows=1, cols=4); _no_border(tq)
    cells = tq.rows[0].cells
    big = lambda t, c: (str(t), 20, True, c)
    cap = lambda t, c: (t, 8.5, True, c)
    cO, cB, cG, cC = RGBColor(0x9A,0x34,0x12), NAVY, RGBColor(0x2E,0x5B,0x1E), RGBColor(0x9A,0x34,0x12)
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell(cells[0], [cap("TIẾN ĐỘ TỔNG KỲ", cO), big(f"{pct_ky}%", cO)], fill=F_ORANGE)
    _set_cell(cells[1], [cap("MỤC TIÊU", cB), big(len(muc_tieu_list), cB)], fill=F_BLUE)
    _set_cell(cells[2], [cap("ĐÃ DUYỆT", cG), big(so_da_duyet, cG)], fill=F_GREEN)
    _set_cell(cells[3], [cap("CẦN SỬA", cC), big(so_can_sua, cC)], fill=F_CORAL)
    for cell in cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Bảng mục tiêu ----
    _heading(doc, "BẢNG MỤC TIÊU VÀ KẾT QUẢ THEN CHỐT", 14, NAVY)
    if not muc_tieu_list:
        doc.add_paragraph("Chưa có mục tiêu nào trong kỳ này.")
    else:
        cols = ["Mục tiêu", "Trạng thái", "Tiến độ", "Kết quả then chốt", "Nhận xét GV"]
        mt_tbl = doc.add_table(rows=1, cols=len(cols)); _no_border(mt_tbl)
        for i, h in enumerate(cols):
            _set_cell(mt_tbl.rows[0].cells[i], [(h, 10.5, True, WHITE)], fill=F_NAVY)
        for mt in muc_tieu_list:
            tt = mt.get("trang_thai", "")
            krs = krs_by_mt.get(mt["id"], [])
            kr_tom = f"{len(krs)} KR" if krs else "Chưa có KR"
            row = mt_tbl.add_row().cells
            _set_cell(row[0], [(mt.get("muc_tieu_lon", ""), 10.5, False, DARK)])
            _set_cell(row[1], [(TRANG_THAI_MT.get(tt, tt), 10.5, False, DARK)], fill=FILL_TRANG_THAI.get(tt))
            _set_cell(row[2], [(f"{mt.get('tien_do_tong') or 0}%", 10.5, False, DARK)], fill=F_ORANGE)
            _set_cell(row[3], [(kr_tom, 10.5, False, DARK)])
            _set_cell(row[4], [(mt.get("nhan_xet_giao_vien") or "—", 10.5, False, DARK)])

        # ---- Chi tiết KR có dữ liệu ----
        co_kr = any(krs_by_mt.get(mt["id"]) for mt in muc_tieu_list)
        if co_kr:
            _heading(doc, "CHI TIẾT KẾT QUẢ THEN CHỐT", 12, DNAVY)
            for mt in muc_tieu_list:
                krs = krs_by_mt.get(mt["id"], [])
                if not krs:
                    continue
                p_mt = doc.add_paragraph()
                p_mt.paragraph_format.space_before = Pt(6)
                rr = p_mt.add_run(f"• {mt.get('muc_tieu_lon', '')}")
                rr.bold = True; rr.font.size = Pt(11); rr.font.color.rgb = DNAVY

                kr_cols = ["KR", "Nội dung", "Khởi điểm → Mục tiêu", "Hiện tại", "Đơn vị", "Tiến độ"]
                kr_tbl = doc.add_table(rows=1, cols=len(kr_cols)); _no_border(kr_tbl)
                for i, h in enumerate(kr_cols):
                    _set_cell(kr_tbl.rows[0].cells[i], [(h, 10, True, WHITE)], fill=F_DNAVY)
                for ki, kr in enumerate(krs, 1):
                    rc = kr_tbl.add_row().cells
                    _set_cell(rc[0], [(str(ki), 10, False, DARK)])
                    _set_cell(rc[1], [(kr.get("noi_dung", ""), 10, False, DARK)])
                    _set_cell(rc[2], [(f"{kr.get('gia_tri_khoi_diem', 0)} → {kr.get('gia_tri_muc_tieu', '')}", 10, False, DARK)])
                    _set_cell(rc[3], [(str(kr.get("gia_tri_hien_tai", 0)), 10, False, DARK)])
                    _set_cell(rc[4], [(kr.get("don_vi", ""), 10, False, DARK)])
                    _set_cell(rc[5], [(f"{kr.get('tien_do_phan_tram', 0)}%", 10, False, DARK)], fill=F_ORANGE)

    # ---- Nhận xét tổng kết GV ----
    if danh_gia and danh_gia.get("nhan_xet_gv"):
        _heading(doc, "NHẬN XÉT TỔNG KẾT CỦA GIÁO VIÊN", 14, NAVY)
        nx = doc.add_table(rows=1, cols=2); _no_border(nx)
        nx.columns[0].width = Cm(4)
        diem = danh_gia.get("diem_so")
        diem_txt = f"{diem} / 5" if diem is not None else "—"
        c0 = nx.rows[0].cells[0]
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell(c0, [("ĐIỂM ĐÁNH GIÁ", 8.5, True, NAVY), (diem_txt, 20, True, NAVY)], fill=F_BLUE)
        for p in c0.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell(nx.rows[0].cells[1], [(danh_gia.get("nhan_xet_gv") or "", 11.5, False, DNAVY)], fill=F_LGRAY, valign="top")

        if danh_gia.get("ky_vong_ky_tiep"):
            p_kv = doc.add_paragraph()
            p_kv.add_run("Kỳ vọng kỳ tiếp: ").bold = True
            p_kv.add_run(danh_gia["ky_vong_ky_tiep"])
        if danh_gia.get("trien_khai_xuat_sac"):
            p_xs = doc.add_paragraph()
            r = p_xs.add_run("✦ Học sinh triển khai xuất sắc — được đề cử khen thưởng")
            r.bold = True; r.font.color.rgb = ORANGE

    # ---- Ý kiến gia đình ----
    if danh_gia and danh_gia.get("phan_hoi_ph"):
        _heading(doc, "Ý KIẾN GIA ĐÌNH", 13, NAVY)
        doc.add_paragraph(danh_gia.get("phan_hoi_ph") or "")

    # ---- Xác nhận ----
    _heading(doc, "XÁC NHẬN", 12, NAVY)
    sg = doc.add_table(rows=2, cols=2); _no_border(sg)
    h0, h1 = sg.rows[0].cells
    _set_cell(h0, [("Giáo viên chủ nhiệm", 11, True, DARK)], fill=F_INFO)
    _set_cell(h1, [("Phụ huynh học sinh", 11, True, DARK)], fill=F_INFO)
    for c in sg.rows[0].cells: c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    b0, b1 = sg.rows[1].cells
    _set_cell(b0, [("\n\n(Ký và ghi rõ họ tên)", 10, False, GRAY)])
    _set_cell(b1, [("\n\n(Ký và ghi rõ họ tên)", 10, False, GRAY)])
    for c in sg.rows[1].cells:
        for p in c.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Footer ----
    p_ft = doc.add_paragraph()
    p_ft.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_ft.add_run(f"Ngày xuất báo cáo: {date.today().strftime('%d/%m/%Y')}")
    r.font.size = Pt(8.5); r.font.color.rgb = GRAY


def _setup_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    return doc


@router.get("/hoc-sinh/{id}")
def xuat_bao_cao_hoc_sinh(id: str, ky_id: str, nguoi_dung=Depends(chi_giao_vien)):
    hs = supabase.table("nguoi_dung").select("*").eq("id", id).execute()
    if not hs.data:
        raise HTTPException(status_code=404, detail="Khong tim thay hoc sinh")

    mt_res = supabase.table("muc_tieu").select("*").eq("hoc_sinh_id", id).eq("ky_danh_gia_id", ky_id)\
        .neq("trang_thai", "nhap").order("ngay_tao").execute()
    dg = supabase.table("danh_gia_cuoi_ky").select("*").eq("hoc_sinh_id", id).eq("ky_danh_gia_id", ky_id).execute()
    ky = supabase.table("ky_danh_gia").select("ten_ky").eq("id", ky_id).execute()
    gv = supabase.table("nguoi_dung").select("ho_ten").eq("id", nguoi_dung["id"]).execute()

    ten_ky = ky.data[0]["ten_ky"] if ky.data else ""
    ten_gv = gv.data[0]["ho_ten"] if gv.data else ""

    krs_by_mt: dict = {}
    for mt in mt_res.data:
        kr_res = supabase.table("ket_qua_then_chot").select("*").eq("muc_tieu_id", mt["id"]).order("thu_tu").execute()
        krs_by_mt[mt["id"]] = kr_res.data

    doc = _setup_doc()
    tao_phieu_hoc_sinh(doc, hs.data[0], mt_res.data, krs_by_mt,
                       dg.data[0] if dg.data else {}, ten_ky, ten_gv)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    ten_file = f"BaoCao_{hs.data[0]['ho_ten'].replace(' ', '_')}_{ten_ky.replace(' ', '_')}.docx"
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _content_disposition(ten_file)})


@router.get("/ca-lop/{ten_lop}")
def xuat_bao_cao_ca_lop(ten_lop: str, ky_id: str, nguoi_dung=Depends(chi_giao_vien)):
    hs_res = supabase.table("nguoi_dung").select("*").eq("ten_lop", ten_lop)\
        .eq("vai_tro", "hoc_sinh").eq("dang_hoat_dong", True).order("so_thu_tu").execute()
    ky = supabase.table("ky_danh_gia").select("ten_ky").eq("id", ky_id).execute()
    gv = supabase.table("nguoi_dung").select("ho_ten").eq("id", nguoi_dung["id"]).execute()

    ten_ky = ky.data[0]["ten_ky"] if ky.data else ""
    ten_gv = gv.data[0]["ho_ten"] if gv.data else ""

    doc = _setup_doc()
    for i, hs in enumerate(hs_res.data):
        if i > 0:
            doc.add_page_break()
        mt_res = supabase.table("muc_tieu").select("*").eq("hoc_sinh_id", hs["id"])\
            .eq("ky_danh_gia_id", ky_id).neq("trang_thai", "nhap").order("ngay_tao").execute()
        dg = supabase.table("danh_gia_cuoi_ky").select("*").eq("hoc_sinh_id", hs["id"])\
            .eq("ky_danh_gia_id", ky_id).execute()

        krs_by_mt: dict = {}
        for mt in mt_res.data:
            kr_res = supabase.table("ket_qua_then_chot").select("*").eq("muc_tieu_id", mt["id"]).order("thu_tu").execute()
            krs_by_mt[mt["id"]] = kr_res.data

        tao_phieu_hoc_sinh(doc, hs, mt_res.data, krs_by_mt,
                           dg.data[0] if dg.data else {}, ten_ky, ten_gv)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    ten_file = f"BaoCao_CaLop_{ten_lop}_{ten_ky.replace(' ', '_')}.docx"
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _content_disposition(ten_file)})
