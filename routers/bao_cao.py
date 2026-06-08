from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import date
from database import supabase
from auth import chi_giao_vien

router = APIRouter()

def tao_phieu_hoc_sinh(doc: Document, hs: dict, muc_tieu_list: list, danh_gia: dict, ten_truong: str, ten_ky: str, ten_gv: str):
    doc.add_heading(ten_truong, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f"PHIEU DANH GIA MUC TIEU — {ten_ky}", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run(f"Ho va ten: {hs['ho_ten']}    Lop: {hs.get('ten_lop', '')}    Ngay xuat: {date.today().strftime('%d/%m/%Y')}")
    info.add_run(f"\nGiao vien chu nhiem: {ten_gv}    Ky danh gia: {ten_ky}")

    doc.add_heading("BANG MUC TIEU", level=3)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"

    headers = ["STT", "Muc tieu lon", "Ket qua then chot", "Chi tieu", "Thuc dat", "Don vi", "Tien do (%)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for idx, mt in enumerate(muc_tieu_list, 1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = mt.get("muc_tieu_lon", "")
        row.cells[2].text = mt.get("ket_qua_then_chot", "")
        row.cells[3].text = str(mt.get("chi_tieu", ""))
        row.cells[4].text = str(mt.get("thuc_dat", 0))
        row.cells[5].text = mt.get("don_vi", "")
        row.cells[6].text = f"{mt.get('tien_do_phan_tram', 0)}%"

    if muc_tieu_list:
        doc.add_heading("NHAN XET CUA GIAO VIEN", level=3)
        nhan_xet_gv = muc_tieu_list[0].get("nhan_xet_giao_vien", "") or ""
        doc.add_paragraph(nhan_xet_gv)

    doc.add_heading("DANH GIA CUA PHU HUYNH", level=3)
    sao = ""
    for mt in muc_tieu_list:
        if mt.get("diem_phu_huynh"):
            sao = "★" * mt["diem_phu_huynh"] + "☆" * (5 - mt["diem_phu_huynh"])
            break
    doc.add_paragraph(f"Danh gia: {sao}" if sao else "Chua co danh gia")

    if danh_gia:
        doc.add_heading("NHAN XET TONG KET CUA GIAO VIEN", level=3)
        doc.add_paragraph(danh_gia.get("nhan_xet_gv") or "")
        doc.add_heading("Y KIEN GIA DINH", level=3)
        doc.add_paragraph(danh_gia.get("phan_hoi_ph") or "")

    doc.add_paragraph()
    ky_ten = doc.add_paragraph()
    ky_ten.add_run("Giao vien chu nhiem" + " " * 30 + "Phu huynh hoc sinh")

@router.get("/hoc-sinh/{id}")
def xuat_bao_cao_hoc_sinh(id: str, ky_id: str, nguoi_dung=Depends(chi_giao_vien)):
    hs = supabase.table("nguoi_dung").select("*").eq("id", id).execute()
    if not hs.data:
        raise HTTPException(status_code=404, detail="Khong tim thay hoc sinh")

    mt = supabase.table("muc_tieu").select("*").eq("hoc_sinh_id", id).eq("ky_danh_gia_id", ky_id).execute()
    dg = supabase.table("danh_gia_cuoi_ky").select("*").eq("hoc_sinh_id", id).eq("ky_danh_gia_id", ky_id).execute()
    ky = supabase.table("ky_danh_gia").select("ten_ky").eq("id", ky_id).execute()
    gv = supabase.table("nguoi_dung").select("ho_ten").eq("id", nguoi_dung["id"]).execute()

    ten_ky = ky.data[0]["ten_ky"] if ky.data else ""
    ten_gv = gv.data[0]["ho_ten"] if gv.data else ""
    ten_truong = "Truong THPT"

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    tao_phieu_hoc_sinh(doc, hs.data[0], mt.data, dg.data[0] if dg.data else {}, ten_truong, ten_ky, ten_gv)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    ten_file = f"BaoCao_{hs.data[0]['ho_ten'].replace(' ', '_')}_{ten_ky.replace(' ', '_')}.docx"
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                              headers={"Content-Disposition": f"attachment; filename={ten_file}"})

@router.get("/ca-lop/{ten_lop}")
def xuat_bao_cao_ca_lop(ten_lop: str, ky_id: str, nguoi_dung=Depends(chi_giao_vien)):
    hs_res = supabase.table("nguoi_dung").select("*").eq("ten_lop", ten_lop).eq("vai_tro", "hoc_sinh").order("ho_ten").execute()
    ky = supabase.table("ky_danh_gia").select("ten_ky").eq("id", ky_id).execute()
    gv = supabase.table("nguoi_dung").select("ho_ten").eq("id", nguoi_dung["id"]).execute()

    ten_ky = ky.data[0]["ten_ky"] if ky.data else ""
    ten_gv = gv.data[0]["ho_ten"] if gv.data else ""
    ten_truong = "Truong THPT"

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    for i, hs in enumerate(hs_res.data):
        if i > 0:
            doc.add_page_break()

        mt = supabase.table("muc_tieu").select("*").eq("hoc_sinh_id", hs["id"]).eq("ky_danh_gia_id", ky_id).execute()
        dg = supabase.table("danh_gia_cuoi_ky").select("*").eq("hoc_sinh_id", hs["id"]).eq("ky_danh_gia_id", ky_id).execute()
        tao_phieu_hoc_sinh(doc, hs, mt.data, dg.data[0] if dg.data else {}, ten_truong, ten_ky, ten_gv)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    ten_file = f"BaoCao_CaLop_{ten_lop}_{ten_ky.replace(' ', '_')}.docx"
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                              headers={"Content-Disposition": f"attachment; filename={ten_file}"})
